"""PDF slide deck -> Word document conversion logic."""
import io
import re
import statistics
import unicodedata
from collections import Counter

import fitz  # PyMuPDF
from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor

BULLET_PREFIXES = ("•", "‣", "▪", "◦", "●", "○", "-", "*", "­", "‐")
# The trailing text is optional: a numbered marker sometimes sits alone on
# its own PDF text line, with the item's actual text on the following line(s).
NUMBERED_RE = re.compile(r"^(\d+[.)]|[a-zA-Z][.)])(?:\s+(.*))?$")
CONTROL_CHAR_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")

MAX_LABEL_IMAGE_HEIGHT_PX = 200  # candidate height for a "text rendered as an image" label
MAX_LABEL_IMAGE_COLORS = 300  # real photos/diagrams have far more distinct colors than a flat text badge

BLOCKLIST_KEYWORDS = ("politecnico", "polimi", "motor system rehabilitation", "semester", "semestre")
ACADEMIC_YEAR_RE = re.compile(r"\b\d{4}\s*/\s*\d{4}\b")
DATE_RE = re.compile(r"\b\d{1,2}[./]\d{1,2}[./]\d{2,4}\b")
EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
SHORT_BLOCK_CHAR_LIMIT = 40  # only auto-drop date-only blocks, not real content that mentions a date

MIN_IMAGE_DIM = 40  # px; skip tiny icons
IMAGE_REPEAT_FRACTION = 0.25  # fraction of the deck's pages an image must recur on to count as a logo
MIN_IMAGE_REPEAT_COUNT = 2  # floor for very short decks
EDGE_BAND_RATIO = 0.08  # top/bottom 8% of the slide = header/footer band
SMALL_FONT_FACTOR = 0.65  # relative to median body font size, for footer detection
H2_FONT_FACTOR = 1.35  # relative to median body font size, for sub-heading detection
H3_FONT_FACTOR = 1.15

LARGE_IMAGE_AREA_RATIO = 0.12  # image covers >=12% of the slide -> treated as a chart/graphic
LARGE_IMAGE_WIDTH_IN = 6.0
LARGE_IMAGE_MAX_HEIGHT_IN = 3.0
SMALL_IMAGE_WIDTH_IN = 2.0
SMALL_IMAGE_MAX_HEIGHT_IN = 1.5

# Slide decks often draw arrows/circles/labels as separate vector shapes on
# top of (or around) a pasted picture -- extracting just the embedded raster
# loses that annotation. These control how "real" annotation shapes are told
# apart from a template's repeated background decoration, and how nearby
# shapes/images get grouped into one rendered snapshot of the whole graphic.
DRAWING_TEMPLATE_FREQUENCY_RATIO = 0.8  # a shape at the same spot on >=80% of pages is template decor
DRAWING_BACKGROUND_AREA_RATIO = 0.5  # a shape covering more than half the slide is a background fill
CLUSTER_MERGE_MARGIN_PT = 20  # merge shapes/images within this many points of each other
MIN_DRAWING_ONLY_CLUSTER_AREA_RATIO = 0.01  # ignore stray single-line clusters with no picture at all
GRAPHIC_RENDER_ZOOM = 2.0  # supersampling factor when rasterizing a clipped slide region

DOCUMENT_FONT = "Aptos"
NARROW_MARGIN_IN = 0.5

# Approximation of Word's built-in "Shaded" design (Design tab > Style Set):
# solid color blocks for the top-level headings, fading to plain colored text
# for the deeper levels.
TITLE_SHADE_FILL = "1F4E79"
TITLE_FONT_COLOR = "FFFFFF"
H1_SHADE_FILL = "2E74B5"
H1_FONT_COLOR = "FFFFFF"
H2_SHADE_FILL = "DEEAF6"
H2_FONT_COLOR = "1F4E79"
H3_FONT_COLOR = "2E74B5"

STYLE_FOR_KIND = {"bullet": "List Bullet", "number": "List Number"}


RIGHT_ARROW_TEXT = "➞"
LEFT_ARROW_TEXT = "<--"


def _arrow_replacement(ch):
    """If ch is an arrow glyph (a named Unicode arrow, or an unmapped
    symbol-font Private Use Area glyph such as the Wingdings arrow
    PowerPoint uses for "leads to" bullets), return its text replacement;
    otherwise return None."""
    cp = ord(ch)
    try:
        name = unicodedata.name(ch)
    except ValueError:
        name = None
    if name and "ARROW" in name:
        is_left = "LEFT" in name and "RIGHT" not in name
        return LEFT_ARROW_TEXT if is_left else RIGHT_ARROW_TEXT
    if 0xE000 <= cp <= 0xF8FF:
        return RIGHT_ARROW_TEXT
    return None


def _convert_arrows(text):
    """Replace any Unicode arrow glyph (any direction, any arrow block) with
    an ASCII/plain-text arrow so it survives into Word text."""
    return "".join(_arrow_replacement(ch) or ch for ch in text)


def _sanitize_text(text):
    """Strip control characters that some fonts/PDF producers leave behind
    for glyphs they can't map (e.g. a missing symbol -> NUL byte); those
    would otherwise crash Word XML generation outright."""
    return CONTROL_CHAR_RE.sub("", text)


def _classify_line_marker(raw_text):
    """Look at a RAW (not yet arrow-converted) line and split off a leading
    list marker. Returns (kind, remainder) where remainder still needs
    _convert_arrows applied. A leading arrow glyph counts as a bullet too --
    some slide templates use an arrow instead of a dot as the bullet mark --
    so it's consumed as the marker instead of becoming inline "➞" text."""
    stripped = raw_text.strip()
    if not stripped:
        return None, stripped
    if stripped[0] in BULLET_PREFIXES:
        return "bullet", stripped[1:].strip()
    m = NUMBERED_RE.match(stripped)
    if m:
        return "number", (m.group(2) or "").strip()
    if _arrow_replacement(stripped[0]) is not None:
        return "bullet", stripped[1:].strip()
    return None, stripped


def _looks_like_page_number(lines):
    return len(lines) == 1 and lines[0].strip().isdigit() and len(lines[0].strip()) <= 3


def _looks_like_institutional_noise(lines):
    joined = " ".join(lines)
    lowered = joined.lower()
    if any(keyword in lowered for keyword in BLOCKLIST_KEYWORDS):
        return True
    if EMAIL_RE.search(joined):
        return True
    if len(joined) <= SHORT_BLOCK_CHAR_LIMIT and (ACADEMIC_YEAR_RE.search(joined) or DATE_RE.search(joined)):
        return True
    return False


def _normalize(text):
    return re.sub(r"\s+", " ", text.strip().lower())


def _block_key(block):
    return _normalize(" ".join(text for text, _kind in block["lines"]))


def _set_run_font(run, name=DOCUMENT_FONT):
    run.font.name = name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), name)


def _set_style_font(document, style_name, name=DOCUMENT_FONT):
    try:
        style = document.styles[style_name]
    except KeyError:
        return
    style.font.name = name
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), name)


def _apply_document_font(document):
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"):
        _set_style_font(document, style_name)


def _shade_paragraph(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _style_heading(paragraph, fill_hex=None, font_hex=None):
    if fill_hex:
        _shade_paragraph(paragraph, fill_hex)
    if font_hex:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(font_hex)


def _set_narrow_margins(document):
    for section in document.sections:
        section.top_margin = Inches(NARROW_MARGIN_IN)
        section.bottom_margin = Inches(NARROW_MARGIN_IN)
        section.left_margin = Inches(NARROW_MARGIN_IN)
        section.right_margin = Inches(NARROW_MARGIN_IN)


def _add_centered_page_number_footer(document):
    for section in document.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        _set_run_font(run)
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = "PAGE"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr_text)
        run._r.append(fld_end)


def _find_list_number_abstract_id(document):
    """Find the abstractNumId backing the built-in "List Number" style, so a
    fresh numbering instance can be created against the same format."""
    numbering_el = document.part.numbering_part.element
    for abstract_num in numbering_el.findall(qn("w:abstractNum")):
        for lvl in abstract_num.findall(qn("w:lvl")):
            p_style = lvl.find(qn("w:pStyle"))
            if p_style is not None and p_style.get(qn("w:val")) == "ListNumber":
                return abstract_num.get(qn("w:abstractNumId"))
    return None


def _new_numbering_instance(document, abstract_num_id):
    """Add a new <w:num> instance referencing the given abstract numbering
    definition and return its numId. Each list gets its own instance so
    Word restarts counting at 1 instead of continuing a shared counter
    across every numbered list in the document."""
    numbering_el = document.part.numbering_part.element
    existing_ids = [int(n.get(qn("w:numId"))) for n in numbering_el.findall(qn("w:num"))]
    new_id = max(existing_ids, default=0) + 1
    num_el = OxmlElement("w:num")
    num_el.set(qn("w:numId"), str(new_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_num_id))
    num_el.append(abstract_ref)
    numbering_el.append(num_el)
    return new_id


def _set_paragraph_num_id(paragraph, num_id):
    pPr = paragraph._p.get_or_add_pPr()
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        numPr = OxmlElement("w:numPr")
        pPr.insert(0, numPr)
    else:
        for child in list(numPr):
            numPr.remove(child)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(num_id_el)


def _extract_page_blocks(page):
    raw = page.get_text("dict")
    blocks = []
    for b in raw.get("blocks", []):
        if b.get("type") != 0:  # 0 = text block
            continue
        lines, sizes, fonts = [], [], []
        for line in b.get("lines", []):
            spans = line.get("spans", [])
            raw_text = _sanitize_text("".join(s.get("text", "") for s in spans).strip())
            if not raw_text:
                continue
            sizes.extend(s.get("size", 0) for s in spans)
            fonts.extend(s.get("font", "") for s in spans)
            kind, remainder = _classify_line_marker(raw_text)
            lines.append((_convert_arrows(remainder), kind))
        plain_lines = [text for text, _kind in lines]
        if not lines or _looks_like_page_number(plain_lines) or _looks_like_institutional_noise(plain_lines):
            continue
        blocks.append({
            "bbox": b["bbox"],
            "lines": lines,
            "avg_size": sum(sizes) / len(sizes) if sizes else 0,
            "all_bold": bool(fonts) and all("bold" in f.lower() for f in fonts),
        })
    blocks.sort(key=lambda blk: (round(blk["bbox"][1], 0), blk["bbox"][0]))
    return blocks


def _merge_lines_into_paragraphs(lines):
    """Merge a flat stream of (text, kind) lines into paragraphs, splitting on
    new bullet/number markers. Takes a plain line list rather than a single
    block's lines because PDF exporters sometimes split one bulleted
    sentence's wrapped continuation into a separate text block -- feeding in
    lines flattened across those block boundaries keeps such sentences whole."""
    paragraphs = []
    current_text, current_kind = None, None
    for text, kind in lines:
        if kind:
            if current_text:
                paragraphs.append((current_text, current_kind))
            current_text, current_kind = text, kind
        elif not current_text:
            # Nothing accumulated yet (or just an empty marker, e.g. a bullet
            # glyph on its own line) -- adopt this text but keep any kind
            # already established by that marker.
            current_text = text
        else:
            current_text += " " + text
    if current_text:
        paragraphs.append((current_text, current_kind))
    return paragraphs


def _looks_like_text_label_image(png_bytes, max_height_px=MAX_LABEL_IMAGE_HEIGHT_PX):
    """Detect a short, near-flat-color image that's really a stylized text
    caption exported as a picture (e.g. WordArt-style headings/percentages
    rendered as a bitmap) rather than a real photo/diagram. Real diagrams
    have far more distinct colors even when small."""
    try:
        with Image.open(io.BytesIO(png_bytes)) as im:
            if im.height > max_height_px:
                return False
            colors = im.convert("RGB").getcolors(maxcolors=MAX_LABEL_IMAGE_COLORS)
        return colors is not None
    except Exception:
        return False


def _drawing_signature(d):
    r = d["rect"]
    return (d.get("type"), round(r.x0), round(r.y0), round(r.x1), round(r.y1))


def _collect_document_data(src):
    """First pass: gather every page's text blocks, image placements, and
    vector-drawing shapes so recurring headers/footers/logos/background
    decoration can be recognized across the whole deck."""
    pages_blocks, pages_images, pages_drawings = [], [], []
    text_freq = Counter()
    xref_freq = Counter()
    digest_freq = Counter()
    drawing_freq = Counter()
    all_sizes = []

    for page in src:
        blocks = _extract_page_blocks(page)
        pages_blocks.append(blocks)
        for b in blocks:
            text_freq[_block_key(b)] += 1
            all_sizes.append(b["avg_size"])

        images = page.get_image_info(hashes=True, xrefs=True)
        pages_images.append(images)
        for info in images:
            if info.get("xref"):
                xref_freq[info["xref"]] += 1
            if info.get("digest"):
                digest_freq[info["digest"]] += 1

        drawings = page.get_drawings()
        pages_drawings.append(drawings)
        for d in drawings:
            drawing_freq[_drawing_signature(d)] += 1

    median_size = statistics.median(all_sizes) if all_sizes else 12
    template_drawing_sigs = {
        s for s, c in drawing_freq.items() if c >= DRAWING_TEMPLATE_FREQUENCY_RATIO * len(src)
    }
    return (
        pages_blocks,
        pages_images,
        pages_drawings,
        text_freq,
        xref_freq,
        digest_freq,
        median_size,
        template_drawing_sigs,
    )


def _filter_text_blocks(blocks, page_height, small_font_threshold, text_freq, page_max_font):
    filtered = []
    for b in blocks:
        in_edge_band = (
            b["bbox"][1] < page_height * EDGE_BAND_RATIO
            or b["bbox"][3] > page_height * (1 - EDGE_BAND_RATIO)
        )
        is_small = b["avg_size"] <= small_font_threshold
        if in_edge_band and is_small:
            continue
        # Anything that repeats verbatim across pages and isn't this page's
        # own (largest-font) title is boilerplate: course name, professor,
        # lesson/date labels, etc. Repetition is checked regardless of font
        # size so it also catches boilerplate set in a normal body-text size.
        is_title_candidate = b["avg_size"] >= page_max_font
        if not is_title_candidate and text_freq[_block_key(b)] >= 2:
            continue
        filtered.append(b)
    return filtered


def _filter_images(images, page_height, page_area, xref_freq, digest_freq, repeat_threshold):
    filtered = []
    for info in images:
        if info.get("width", 0) < MIN_IMAGE_DIM or info.get("height", 0) < MIN_IMAGE_DIM:
            continue
        bbox = info["bbox"]
        # Only treat as a corner logo/watermark if it sits ENTIRELY within the
        # top/bottom margin band; a large chart merely extending past that
        # line (but mostly in the body) must survive.
        fully_in_top_band = bbox[3] <= page_height * EDGE_BAND_RATIO
        fully_in_bottom_band = bbox[1] >= page_height * (1 - EDGE_BAND_RATIO)
        if fully_in_top_band or fully_in_bottom_band:
            continue
        # A logo repeats on nearly every page; a reference diagram the
        # lecturer reuses across a handful of slides is legitimate content,
        # not decoration, so only the former should be filtered out.
        if info.get("xref") and xref_freq[info["xref"]] >= repeat_threshold:
            continue
        if info.get("digest") and digest_freq[info["digest"]] >= repeat_threshold:
            continue
        filtered.append(info)
    return filtered


def _page_annotation_rects(drawings, template_drawing_sigs, page_area):
    """Vector shapes drawn directly on the slide (arrows, circles, callout
    lines) that aren't part of the template's repeated background decor."""
    rects = []
    for d in drawings:
        if _drawing_signature(d) in template_drawing_sigs:
            continue
        r = d["rect"]
        if r.width <= 0 or r.height <= 0:
            continue
        if r.width * r.height >= DRAWING_BACKGROUND_AREA_RATIO * page_area:
            continue
        rects.append(fitz.Rect(r))
    return rects


def _cluster_rects(items, margin):
    """Group (rect, label) entries whose boxes overlap or sit within `margin`
    points of each other, so a pasted picture and the arrows/circles drawn
    around it become one combined region. Returns a list of
    {"rect": combined_rect, "labels": set(...)}."""
    n = len(items)
    if n == 0:
        return []
    parent = list(range(n))

    def find(x):
        while parent[x] != x:
            parent[x] = parent[parent[x]]
            x = parent[x]
        return x

    def union(a, b):
        ra, rb = find(a), find(b)
        if ra != rb:
            parent[ra] = rb

    changed = True
    while changed:
        changed = False
        roots = {}
        for i in range(n):
            root = find(i)
            if root not in roots:
                roots[root] = fitz.Rect(items[i][0])
            else:
                roots[root] |= items[i][0]
        root_ids = list(roots.keys())
        for a in range(len(root_ids)):
            for b in range(a + 1, len(root_ids)):
                if find(root_ids[a]) == find(root_ids[b]):
                    continue
                rect_a, rect_b = roots[root_ids[a]], roots[root_ids[b]]
                expanded = fitz.Rect(rect_a.x0 - margin, rect_a.y0 - margin, rect_a.x1 + margin, rect_a.y1 + margin)
                if expanded.intersects(rect_b):
                    union(root_ids[a], root_ids[b])
                    changed = True

    clusters = {}
    for i in range(n):
        root = find(i)
        if root not in clusters:
            clusters[root] = {"rect": fitz.Rect(items[i][0]), "labels": set()}
        else:
            clusters[root]["rect"] |= items[i][0]
        clusters[root]["labels"].add(items[i][1])
    return list(clusters.values())


def _render_clip_png(page, rect, zoom=GRAPHIC_RENDER_ZOOM):
    matrix = fitz.Matrix(zoom, zoom)
    pixmap = page.get_pixmap(clip=rect, matrix=matrix)
    return pixmap.tobytes("png")


def _detect_title(blocks, page_height):
    if not blocks:
        return None, blocks
    max_size = max(b["avg_size"] for b in blocks)
    for b in blocks:
        if b["avg_size"] == max_size and b["bbox"][1] < page_height * 0.35:
            title = " ".join(text for text, _kind in b["lines"])
            body_blocks = [x for x in blocks if x is not b]
            return title, body_blocks
    return None, blocks


def _looks_like_cover_slide_label(block):
    """Short, non-list fragments like 'Lesson', 'of', a lecture topic, or a
    professor's name, typically laid out as separate boxes on a title
    slide rather than real sentences."""
    if block["lines"][0][1]:  # starts with a bullet/number marker
        return False
    word_count = sum(len(text.split()) for text, _kind in block["lines"])
    return len(block["lines"]) <= 2 and word_count <= 4


def _classify_subheading(block, median_size):
    """A short, single-line block that is noticeably larger than normal body
    text (but isn't the slide title) is treated as a Heading 2/3. A short,
    entirely bold line at ordinary body size (e.g. "Pros"/"Cons",
    "Advantages") is also a sub-header even without a size difference."""
    if len(block["lines"]) != 1 or block["lines"][0][1]:
        return None
    text = block["lines"][0][0]
    if block["avg_size"] >= median_size * H2_FONT_FACTOR:
        return "heading2"
    if block["avg_size"] >= median_size * H3_FONT_FACTOR:
        return "heading3"
    if block.get("all_bold") and len(text.split()) <= 4:
        return "heading3"
    return None


def _add_sized_picture(document, png_bytes, target_width_in, max_height_in):
    width_kwargs = {"width": Inches(target_width_in)}
    try:
        with Image.open(io.BytesIO(png_bytes)) as im:
            px_w, px_h = im.size
        if px_w and px_h:
            projected_height_in = target_width_in * (px_h / px_w)
            if projected_height_in > max_height_in:
                width_kwargs = {"height": Inches(max_height_in)}
    except Exception:
        pass
    document.add_picture(io.BytesIO(png_bytes), **width_kwargs)
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def extract_content(pdf_bytes, title="Presentazione"):
    """Parse the PDF into an ordered list of content items (title, headings,
    paragraphs, images) independent of how the Word document gets built, so
    the same extraction can feed both the original and a translated build."""
    src = fitz.open(stream=pdf_bytes, filetype="pdf")
    (
        pages_blocks,
        pages_images,
        pages_drawings,
        text_freq,
        xref_freq,
        digest_freq,
        median_size,
        template_drawing_sigs,
    ) = _collect_document_data(src)
    small_font_threshold = median_size * SMALL_FONT_FACTOR

    items = [{"kind": "title", "text": title}]
    prev_title_norm = None
    seen_in_section = set()
    image_repeat_threshold = max(MIN_IMAGE_REPEAT_COUNT, IMAGE_REPEAT_FRACTION * len(src))

    for page_index, page in enumerate(src):
        page_height = page.rect.height
        page_area = page.rect.width * page_height
        raw_blocks = pages_blocks[page_index]
        page_max_font = max((b["avg_size"] for b in raw_blocks), default=0)

        blocks = _filter_text_blocks(raw_blocks, page_height, small_font_threshold, text_freq, page_max_font)
        images = _filter_images(
            pages_images[page_index], page_height, page_area, xref_freq, digest_freq, image_repeat_threshold
        )

        slide_title, body_blocks = _detect_title(blocks, page_height)

        is_first_page = page_index == 0
        if is_first_page:
            # A cover slide's stray labels (lesson number, topic, professor
            # name) aren't real content. Only drop the page's own "title"
            # too when NOTHING real is left once those are removed -- a
            # short but genuine slide title (e.g. "Analisi della stagione")
            # must survive when it's followed by real body content.
            filtered_body = [b for b in body_blocks if not _looks_like_cover_slide_label(b)]
            if not filtered_body and slide_title and _looks_like_cover_slide_label({"lines": [(slide_title, None)]}):
                slide_title = None
            body_blocks = filtered_body

        title_norm = _normalize(slide_title) if slide_title else None

        if title_norm and title_norm != prev_title_norm:
            items.append({"kind": "heading1", "text": slide_title})
            prev_title_norm = title_norm
            seen_in_section = set()

        pending_text, pending_kind = None, None

        def _flush_pending():
            nonlocal pending_text, pending_kind
            if pending_text:
                norm = _normalize(pending_text)
                if norm not in seen_in_section:
                    seen_in_section.add(norm)
                    items.append({"kind": "paragraph", "text": pending_text, "list_style": pending_kind})
            pending_text, pending_kind = None, None

        for b in body_blocks:
            subheading_kind = _classify_subheading(b, median_size)
            if subheading_kind:
                _flush_pending()
                subheading_text = b["lines"][0][0]
                text_norm = _normalize(subheading_text)
                if text_norm not in seen_in_section:
                    seen_in_section.add(text_norm)
                    items.append({"kind": subheading_kind, "text": subheading_text})
                continue

            block_paragraphs = _merge_lines_into_paragraphs(b["lines"])
            if not block_paragraphs:
                continue

            first_text, first_kind = block_paragraphs[0]
            rest = block_paragraphs[1:]
            # A bulleted/numbered sentence's wrapped continuation sometimes
            # lands in a separate PDF text block from its own marker (no
            # marker of its own). Only bridge that gap when the still-open
            # item actually had a marker -- two adjacent PLAIN blocks (e.g.
            # independent diagram captions/labels) must stay separate.
            if pending_text is not None and pending_kind and not first_kind:
                pending_text += " " + first_text
            else:
                _flush_pending()
                pending_text, pending_kind = first_text, first_kind

            for text, kind in rest:
                _flush_pending()
                pending_text, pending_kind = text, kind

        _flush_pending()

        # A pasted picture's arrows/circles/callout labels are often drawn as
        # separate vector shapes on top of or around it, not baked into the
        # raster -- cluster each image with any nearby annotation shapes and
        # render the WHOLE region together so nothing gets lost.
        annotation_rects = _page_annotation_rects(pages_drawings[page_index], template_drawing_sigs, page_area)
        cluster_items = [(fitz.Rect(info["bbox"]), "image") for info in images]
        cluster_items += [(r, "drawing") for r in annotation_rects]
        clusters = _cluster_rects(cluster_items, CLUSTER_MERGE_MARGIN_PT)

        for cluster in clusters:
            rect = cluster["rect"]
            has_image = "image" in cluster["labels"]
            has_annotation = "drawing" in cluster["labels"]
            area_ratio = (rect.width * rect.height) / page_area

            if not has_image and area_ratio < MIN_DRAWING_ONLY_CLUSTER_AREA_RATIO:
                # Pure vector shapes with no picture and too small to be a
                # real diagram (e.g. a single stray line) -- skip.
                continue

            try:
                png_bytes = _render_clip_png(page, rect)
            except Exception:
                continue

            if has_image and not has_annotation:
                # Behaves like a single plain picture (nothing merged in) --
                # the WordArt-style flat-label filter still applies.
                label_height_threshold = MAX_LABEL_IMAGE_HEIGHT_PX * GRAPHIC_RENDER_ZOOM
                if _looks_like_text_label_image(png_bytes, label_height_threshold):
                    continue

            is_large = area_ratio >= LARGE_IMAGE_AREA_RATIO
            if is_first_page and not is_large:
                # Cover slides carry logos/decorative badges, not content;
                # only a genuinely large graphic there is worth keeping.
                continue
            if is_large:
                target_width, max_height = LARGE_IMAGE_WIDTH_IN, LARGE_IMAGE_MAX_HEIGHT_IN
            else:
                target_width, max_height = SMALL_IMAGE_WIDTH_IN, SMALL_IMAGE_MAX_HEIGHT_IN
            items.append({
                "kind": "image",
                "png_bytes": png_bytes,
                "target_width": target_width,
                "max_height": max_height,
            })

    return items


_HEADING_STYLE = {
    "title": (0, TITLE_SHADE_FILL, TITLE_FONT_COLOR),
    "heading1": (1, H1_SHADE_FILL, H1_FONT_COLOR),
    "heading2": (2, H2_SHADE_FILL, H2_FONT_COLOR),
    "heading3": (3, None, H3_FONT_COLOR),
}


def build_document(items):
    """Render a list of content items (see extract_content) into a styled
    Word document and return its bytes."""
    document = Document()
    _apply_document_font(document)
    _set_narrow_margins(document)
    _add_centered_page_number_footer(document)

    list_number_abstract_id = _find_list_number_abstract_id(document)
    current_number_list_id = None

    for item in items:
        kind = item["kind"]
        if kind in _HEADING_STYLE:
            level, fill_hex, font_hex = _HEADING_STYLE[kind]
            paragraph = document.add_heading(item["text"], level=level)
            _style_heading(paragraph, fill_hex, font_hex)
            current_number_list_id = None
        elif kind == "paragraph":
            list_style = item.get("list_style")
            style = STYLE_FOR_KIND.get(list_style)
            paragraph = document.add_paragraph(item["text"], style=style)
            for run in paragraph.runs:
                _set_run_font(run)
            if list_style == "number":
                if current_number_list_id is None and list_number_abstract_id is not None:
                    current_number_list_id = _new_numbering_instance(document, list_number_abstract_id)
                if current_number_list_id is not None:
                    _set_paragraph_num_id(paragraph, current_number_list_id)
            else:
                current_number_list_id = None
        elif kind == "image":
            try:
                _add_sized_picture(document, item["png_bytes"], item["target_width"], item["max_height"])
            except Exception:
                continue
            current_number_list_id = None

    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


TRANSLATION_ERROR_MARKERS = (
    "error 500",
    "server error",
    "that's an error",
    "that’s an error",
    "that's all we know",
    "that’s all we know",
)
TRANSLATION_CHUNK_SIZE = 25
TRANSLATION_CHUNK_DELAY_SEC = 1.0
TRANSLATION_ITEM_DELAY_SEC = 0.3


def _is_valid_translation(result):
    if not result or not result.strip():
        return False
    lowered = result.lower()
    return not any(marker in lowered for marker in TRANSLATION_ERROR_MARKERS)


def translate_items(items, target_lang="it"):
    """Return a copy of items with every text field translated; falls back
    to the original text for anything that can't be translated. The free
    Google Translate endpoint occasionally returns its own HTML error page
    body as if it were a translation (no exception raised) when it's rate
    limited, so every result is validated before being accepted -- a
    rejected/failed item just keeps its original text rather than being
    replaced with garbage."""
    import time

    from deep_translator import GoogleTranslator

    translator = GoogleTranslator(source="auto", target=target_lang)
    text_indices = [i for i, item in enumerate(items) if item["kind"] != "image" and item["text"].strip()]
    texts = [items[i]["text"] for i in text_indices]

    new_items = [dict(item) for item in items]
    if not texts:
        return new_items

    for chunk_start in range(0, len(texts), TRANSLATION_CHUNK_SIZE):
        chunk_indices = text_indices[chunk_start:chunk_start + TRANSLATION_CHUNK_SIZE]
        chunk_texts = texts[chunk_start:chunk_start + TRANSLATION_CHUNK_SIZE]

        translated = None
        try:
            translated = translator.translate_batch(chunk_texts)
        except Exception:
            translated = None

        if (
            translated
            and len(translated) == len(chunk_texts)
            and all(_is_valid_translation(t) for t in translated)
        ):
            for idx, translated_text in zip(chunk_indices, translated):
                new_items[idx]["text"] = translated_text
        else:
            # Batch failed or looked corrupted -- retry this chunk one item
            # at a time, validating each and keeping the original on failure.
            for idx in chunk_indices:
                try:
                    result = translator.translate(new_items[idx]["text"])
                except Exception:
                    result = None
                if _is_valid_translation(result):
                    new_items[idx]["text"] = result
                time.sleep(TRANSLATION_ITEM_DELAY_SEC)

        time.sleep(TRANSLATION_CHUNK_DELAY_SEC)

    return new_items


def pdf_to_docx(pdf_bytes, title="Presentazione"):
    return build_document(extract_content(pdf_bytes, title))


def pdf_to_docx_italian(pdf_bytes, title="Presentazione"):
    items = extract_content(pdf_bytes, title)
    return build_document(translate_items(items, target_lang="it"))
